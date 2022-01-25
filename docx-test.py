import re, os, openpyxl as oxl
from time import sleep
from docx import Document
from datetime import date as dt
import inquirer
from tqdm import tqdm
# Clearance log
import firebase_admin
from firebase_admin import credentials
from firebase_admin import firestore

cred = credentials.Certificate("./sa.json")
firebase_admin.initialize_app(cred)
db = firestore.client()

# from phaxio import PhaxioApi

# convert('./medrecs/faxcover.docx')

# Search the doc(s) and replace data
def docx_replace_regex(doc_obj, regex , replace):
    
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

# print(patient['yourName'])
def main():
    os.system('cls')
    import PyPDF2 as pdf

    doc_question = [inquirer.Checkbox('docs', message=f'What documents do you need?', choices=['Anticoagulant', 'A1c Tests', 'Pacemaker', 'Clearance'])]
    prompt = inquirer.prompt(doc_question)
    print(prompt['docs'])
    raw_docs = []
    docs = [Document('./medrecs/faxcover.docx')]

    # Take choices and add them to raw_docs
    for choice in prompt['docs']:
        raw_docs.append(choice)

    # print(f"The docs are: {docs}")
    for doc in raw_docs:
        docs.insert(0, Document(f'./medrecs/{doc}.docx'))
        # print(doc)
    # print(f"The new docs are: {docs}")
    prompt['docs'].append('Faxcover')
    patient = {
    "ptName": input(f'What is the name of the patient?\n'),
    "dateOfBirth": input(f"What is the patient's date of birth?\n"),
    "procedureDate": input(f'What is the date of the procedure?\n'),
    "procedureName": input(f'What is the procedure name? (i.e. PVP (Photovaporization of the prostate))\n'),
    "anesthesiologistName": input(f'Who is the surgeon? (Kaplan, Wong...)\n'),
    "drName": input(f'What is the name of the doctor you are contacting? (Name only! No "Dr." needed!)\n'),
    "pNumber": input(f'What is the phone number of the facility you are faxing?\n'),
    "fNumber": input(f'What is the number you are faxing to?\n'),
    "forms": prompt["docs"],
    "dateOfFax": dt.today().strftime("%B %d, %Y"),
    "numberOfPages": str(len(docs)),
    "urgency": 'URGENT',
    "yourName": 'Zackery H./Franchesca G.',
    }
    os.system('cls')
    # Input data into tables
    print('Applying information to templates...')
    for doc in docs:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word, replacement in patient.items():
                            word_re=re.compile(word)
                            docx_replace_regex(doc, word_re , replacement)
    # Input data into paragraphs
    for doc in docs:
        for word, replacement in patient.items():
            word_re=re.compile(word)
            docx_replace_regex(doc, word_re , replacement)
    # Save the changed files "as"
    for i in tqdm(range(50), desc="Applying info to templates", ascii=True):
        sleep(0.01)
    x = 0
    for doc in docs:
        doc.save(f"file_{x}.docx")
            # doc.save(f"pacemaker_{patient['ptName']}.docx")
        if docs.index(doc) != len(docs) - 1:
            x += 1
            
    # Edit files to be PDFs
    print('Converting to PDFs...')
    os.system('powershell docx2pdf .')
    sleep(2)
    mergeFile = pdf.PdfFileMerger()
    pdfFiles = []
    while x >= 0:
        pdfFiles.append(pdf.PdfFileReader(f"file_{x}.pdf", 'rb'))
        x -= 1

    for pdf in pdfFiles:
        mergeFile.append(pdf)

    names = patient['ptName'].split(' ')
    os.system('powershell rm *.docx')
    os.system('powershell rm *.pdf')
    mergeFile.write(f'Request- {names[1]}, {names[0]} Med Recs Req.pdf')
    
    # Write to excel
    wb = oxl.Workbook() # Create workbook
    ws = wb.active # Set active worksheet

    col = 1 
    igs = ['urgency', 'yourName', 'numberOfPages', 'fNumber']
    patient['forms'] = ', '.join(patient['forms'])
    patient['pNumber'] += '/' + patient['fNumber']
    patient['pNumber'] = patient['pNumber'].replace(' ', '.')
    patient['pNumber'] = patient['pNumber'].replace('-', '.')
    for key, value in patient.items():
        try:
            igs.index(key) # See if key is ignored
        except ValueError:
            ws.cell(1, col, value) # Write the data
            col += 1
        pass
    wb.save(f'Request- {names[1]}, {names[0]} Med Recs Req.xlsx')

    print(f'Adding {patient["ptName"]}...')
    os.system(r"powershell mv *.pdf 'S:\'")
    # db.collection('Patient Clearance').document(f"{patient['ptName']} ({patient['dateOfBirth']})").set(patient)


    print('Done!')
    sleep(1)
    os.system('cls')
    return True
main()








 # def faxIt(pt):
    #     fileFaxing = f"Request- {names[1]}, {names[0]} med recs req.pdf"
    #     phaxio = PhaxioApi('326hm7hwwjgtnh18qo66lauwmx4oc8wf2vbth6he', '014a3atyonq17kx4ftixk6orfs0addix5yowk9jd')
    #     phaxio.Fax.send(
    #         to='6233221504',
    #         files=fileFaxing
    #     )
    # faxIt(patient)