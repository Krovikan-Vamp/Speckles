import re
import os
import subprocess
import sys
import inquirer
import firebase_admin
import openpyxl as oxl
import rsa
import requests
import json
import random
from time import sleep
from docx import Document
from alive_progress import alive_bar
from datetime import date as dt
from colorama import Fore, Style
# from phaxio import PhaxioApi
# Fax Information
from selenium import webdriver
import chromedriver_autoinstaller
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
import pyautogui as pag
# Firebase and auto_suggestion
from firebase_admin import credentials, firestore
from prompt_toolkit import prompt, HTML
from prompt_toolkit.completion import WordCompleter, FuzzyCompleter

# Search the doc(s) and replace data


def docx_replace_regex(doc_obj, regex, replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def decrypt(data):
    arr = []
    for char in data:
        temp = int(char) - 25
        arr.append(chr(temp))
        # print(char)

    return ''.join(arr)


def encrypt2(data):
    def split(word):
        return [char for char in word]

    arr = []
    chars = split(data)

    for char in chars:
        arr.append(ord(char) + 25)

    return arr


def speckles():
    global responses
    weather_speck = json.loads(requests.get(
        r'https://api.openweathermap.org/data/2.5/weather?q=Sun%20City&units=imperial&appid=396b8dda92a5079f3bbf2704d32fc382').text)

    yeQuotes = []
    bookQuotes = []
    # activities = []

    bQuote = json.loads(requests.get(
        'https://www.quotepub.com/api/widget/?type=rand&limit=5').text)
    for n in range(5):
        # Kayne quotes
        quote = json.loads(requests.get('https://api.kanye.rest/').text)
        yeQuotes.append(quote['quote'])

        # Activities
        # action = json.loads(requests.get(
        #     'https://www.boredapi.com/api/activity').text)
        # activities.append(
        #     f"With {action['participants']} person/people you can... {action['activity']}.")
        bookQuotes.append(
            f'"{bQuote[n]["quote_body"]}" - {bQuote[n]["quote_author"]}')

    responses = [{
        'name': 'Speckles, the Meteorologist 📡',
        'prompts': ['<p style="text-align: center; padding: 10px;">🐍 - "Good morning, welcome to the Speckles weather channel!" 😎   </p> ', f'<p style="text-align: center; padding: 10px;">Today is 📅 {dt.today().strftime("%B %d, %Y")} 🐍 - "When is the next holiday again?"</p>', f"<p style='text-align: center; padding: 10px;'>It's feeling like {weather_speck['main']['feels_like']}°F 🐍 - 'Always too damn hot'</p> ", f"The clouds cover {weather_speck['clouds']['all']}% of the sky ☁️ 🐍 - 'Wish it was more...'", f"Get ready for a daily high of: {weather_speck['main']['temp_max']}°F 🤠 'Yee haw'"]
    }, {
        'name': 'Speckles, the Kanye Enthusiast 🎵 ',
        'prompts': yeQuotes
    }, {
        'name': 'Speckles, The Philosopher 🤓',
        'prompts': bookQuotes
    }]


def decryptRSA(key, encrypted_data):
    new_data = rsa.decrypt(encrypted_data, key)
    return new_data.decode()


def decryptRSA(key, encrypted_data):
    new_data = rsa.decrypt(encrypted_data, key)
    return new_data.decode()


def main():
    speck_time = random.choice(responses)
    privKey = rsa.PrivateKey.load_pkcs1(open("privKey.pem", "rb").read())
    os.system('cls')
    firebase_admin.initialize_app(credentials.Certificate(
        './sa.json'), {'storageBucket': 'fourpeaks-sc.appspot.com'})
    db = firestore.client()

    # Create Suggestions
    raw_docs = db.collection(u'Auto Suggestions').stream()
    docs = []
    name_suggs = []
    suggestion_list = {'fax': [], 'phone': [], 'dr': [], 'procedure': [], 'surgeons': [
        'LaTowsky', 'Mai', 'Kaplan', 'Kundavaram', 'Stern', 'Klauschie', 'Schlaifer', 'Jones', 'Wong', 'Devakumar']}

    raw_data = db.collection('Names Collected').stream()

    for doc in raw_data:
        file = doc.to_dict()
        stuff = decryptRSA(privKey, file["data"])

        try:
            name_suggs.index(stuff)
        except ValueError:
            name_suggs.append(stuff)
            pass
    os.system('cls')

    # Makes the docs usable as dicts
    with alive_bar(len(docs), title='Creating suggestions', theme='classic') as bar:
        for doc in raw_docs:
            docs.append(doc.to_dict())
            # sleep(0.01)
        for doc in docs:
            # sleep(.25)
            for [key, value] in doc.items():
                doc[key] = decrypt(value)
        keys = ['fax', 'phone', 'dr', 'procedure']

        for doc in docs:
            for key in keys:
                # Append suggestion to list if it doesn't already exitst
                try:
                    suggestion_list[key].index(doc[key])
                except ValueError:
                    suggestion_list[key].append(doc[key])
                    pass
            # sleep(0.5)
            bar()

    os.system('cls')
    import PyPDF2 as pdf  # This module does not like to import at the top 😖

    print(f'You got {speck_time["name"]}!\n')
    doc_question = [inquirer.Checkbox('docs', message=f'What documents do you need? 📝 ', choices=[
                                      'Anticoagulant', 'A1c Tests', 'Pacemaker', 'Clearance'])]
    specialist_question = [inquirer.List(
        'specialist', message='Which specialist are you contacting? 🩺 ', choices=['Cardio', 'PCP', 'Endo', 'Pulmo', 'Onco', 'Nephro', 'Neuro', 'Hema'])]
    og_prompt = inquirer.prompt(doc_question)
    spec_prompt = inquirer.prompt(specialist_question)

    raw_docs = []
    docs = [Document('./medrecs/faxcover.docx')]

    # Take choices and add them to raw_docs
    for choice in og_prompt['docs']:
        raw_docs.append(choice)

    for doc in raw_docs:
        docs.insert(0, Document(f'./medrecs/{doc}.docx'))

    og_prompt['docs'].append('Faxcover')
    patient = {
        "ptName": prompt(f'What is the name of the patient? 🧍\n', completer=FuzzyCompleter(WordCompleter(name_suggs)), complete_in_thread=True, complete_while_typing=True),
        "dateOfBirth": input(f"What is the patient's {Fore.RED}date of birth{Style.RESET_ALL}? 📅\n"),
        "procedureDate": input(f'What is the {Fore.RED}date of the procedure{Style.RESET_ALL}? 📅\n'),
        "sentBy": "ZH",  # If you aren't Zack, please adjust to your initials!
        "procedureName": prompt(f'What is the procedure name? (i.e. PVP (Photovaporization of the prostate)) 🔪\n', bottom_toolbar=HTML(speck_time["prompts"][0]), completer=FuzzyCompleter(WordCompleter(suggestion_list['procedure'])), complete_in_thread=True, complete_while_typing=True),
        "anesthesiologistName": prompt(f'Who is the surgeon? (Kaplan, Wong...) 🩺\n', bottom_toolbar=HTML(speck_time["prompts"][1]), completer=FuzzyCompleter(WordCompleter(suggestion_list['surgeons'])), complete_in_thread=True, complete_while_typing=True),
        "drName": prompt(f'What is the name of the doctor you are contacting? (Name only! No "Dr." needed!)\n', bottom_toolbar=HTML(speck_time["prompts"][2]), completer=FuzzyCompleter(WordCompleter(suggestion_list['dr'])), complete_in_thread=True, complete_while_typing=True),
        "pNumber": prompt(f'What is the phone number of the facility you are faxing? 📞\n', bottom_toolbar=HTML(speck_time["prompts"][3]), completer=FuzzyCompleter(WordCompleter(suggestion_list['phone'])), complete_in_thread=True, complete_while_typing=True),
        "fNumber": prompt(f'What is the number you are faxing to? 📠\n', bottom_toolbar=HTML(speck_time["prompts"][4]), completer=FuzzyCompleter(WordCompleter(suggestion_list['fax'])), complete_in_thread=True, complete_while_typing=True),
        "forms": og_prompt["docs"],
        "dateOfFax": dt.today().strftime("%B %d, %Y"),
        "numberOfPages": str(len(docs)),
        "urgency": '- STAT',
        "yourName": 'Zackery H.',
    }
    os.system('cls')

    # Input data into tables
    for doc in docs:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word, replacement in patient.items():
                            word_re = re.compile(word)
                            docx_replace_regex(doc, word_re, replacement)

    # Input data into paragraphs
    for doc in docs:
        for word, replacement in patient.items():
            word_re = re.compile(word)
            docx_replace_regex(doc, word_re, replacement)

    # Save the changed files "as"
    x = 0
    with alive_bar(title='Applying information to templates', bar='fish') as bar:
        for doc in docs:
            doc.save(f"file_{x}.docx")
            # doc.save(f"pacemaker_{patient['ptName']}.docx")
            if docs.index(doc) != len(docs) - 1:
                x += 1
            bar()

    # Edit files to be PDFs
    print('Converting to PDFs...')
    os.system('powershell docx2pdf .')
    sleep(2)
    mergeFile = pdf.PdfFileMerger()
    pdfFiles = []
    while x >= 0:
        pdfFiles.append(pdf.PdfFileReader(f"file_{x}.pdf", 'rb'))
        x -= 1

    for pdf in pdfFiles:
        mergeFile.append(pdf)

    names = patient['ptName'].split(' ')
    os.system('powershell rm *.docx')
    os.system('powershell rm *.pdf')
    mergeFile.write(
        f'Request- {names[1]}, {names[0]} {spec_prompt["specialist"]} Med Recs Req.pdf')

    # Add to 'suggestion' collections
    new_info = {'fax': encrypt2(patient['fNumber']), 'phone': encrypt2(patient['pNumber']), 'drType': encrypt2(spec_prompt['specialist']),
                'dr': encrypt2(patient['drName']), 'procedure': encrypt2(patient['procedureName']), 'n': encrypt2(patient['ptName'])}
    db.collection('Auto Suggestions').document().set(new_info)

    encName = rsa.encrypt(patient['ptName'].encode(), pubKey)
    db.collection('Names Collected').document().set({'data': encName})

    # Write to excel
    wb = oxl.load_workbook('Requests.xlsx')  # Load records requests workbook
    ws = wb.active  # Set active worksheet

    col = 1
    igs = ['urgency', 'yourName', 'numberOfPages', 'fNumber']
    patient['forms'] = ', '.join(patient['forms'])
    patient['pNumber'] += '/' + patient['fNumber']
    patient['pNumber'] = patient['pNumber'].replace(' ', '.')
    patient['pNumber'] = patient['pNumber'].replace('-', '.')
    patient['drName'] = f'{spec_prompt["specialist"]}: {patient["drName"]}'

    rowToWrite = ws.max_row + 1
    for key, value in patient.items():
        try:
            igs.index(key)  # See if key is ignored
        except ValueError:
            ws.cell(rowToWrite, col, value)  # Write the data
            col += 1
        pass
    wb.save(f'Requests.xlsx')

    # Move patient to the scans folder

    p = subprocess.Popen(
        ['powershell.exe', '.\\file_mgmt.ps1'], stdout=sys.stdout)
    sleep(1)

    def fax(patient):
        faxNo = patient["fNumber"].replace('.', '')
        # forms = ', '.join(patient["forms"])
        chromedriver_autoinstaller.install()
        driver = webdriver.Chrome(service=Service())
        driver.get('https://secure.ipfax.net/')
        sleep(2)
        # Login to the fax service
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div[1]/input').send_keys('6233221504')
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div[2]/input').send_keys('6233221504')
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div[4]/input').click()
        sleep(2)
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div[5]/div/input').click()

        #  Fill out the fields
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[1]/div[1]/div/input').send_keys('Medical Records')
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[1]/div[2]/div/input').send_keys(f'{patient["ptName"]} -- {patient["forms"]}')
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[2]/div[1]/div[1]/input').send_keys(faxNo)
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[2]/div[1]/div[2]/input').click()
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[2]/div[7]/button').click()
        sleep(1)
        # Choose file with PyAutoGui
        pag.hotkey('ctrl', 'l')
        pag.write('S:\\')
        pag.press('enter')
        sleep(10)

        # Send the fax
        driver.find_element(
            by=By.XPATH, value='/html/body/form[1]/div/div/div[2]/div[3]/div[1]/input').click()
        sleep(2)
        driver.close()

    fax(patient)
    print("Don't forget to add them to the log! Check the new '.xlsx' file!")
    sleep(5)
    os.system('cls')
    return True


# Load public key from dir
pubKey = rsa.PublicKey.load_pkcs1(open("pubKey.pem", "rb").read())
speckles()
try:
    main()
except KeyboardInterrupt:
    print('Program killed 🔪 X_X 🔪')

"""
def faxIt(pt):
    fileFaxing = f"Request- {names[1]}, {names[0]} {spec_prompt['specialist']} Med Recs Req.pdf"
    pt['fNumber'] = pt['fNumber'].replace(' ', '.')
    pt['fNumber'] = pt['fNumber'].replace('-', '.')
    faxNumber = pt['fNumber']
    phaxio = PhaxioApi('7kgdmam2fdb4b3526751we7s5dim88s0u7n3kwxe',
                       'bncxks0w76uods89k57h2kjvo8ykjxlytd306vnp')
    phaxio.Fax.send(
        to=faxNumber,
        files=fileFaxing,
        # direction='received'
    )
faxIt(patient)
"""
